"""
This file is part of JobMine EZSearch.

JobMine EZSearch is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

JobMine EZSearch is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with JobMine EZSearch.  If not, see <http://www.gnu.org/licenses/>.
"""

from flask import Blueprint, render_template, jsonify
import logic
from docx import Document
from docx.shared import Pt
import db #this is used by create_word_document. SHould place this elsewhere.

#will need to refactor and place the word creating code somewhere else.


downloads = Blueprint('downloads', __name__)

@downloads.route('/exportHtml/<int:offset>')
def export_html(offset):
    """Displays a HTML page containing 1000 jobs. Offset determines which 1000"""

    jobs = helper_download('HTML',offset)
    return render_template("ExportHTML.html", jobs = jobs)



@downloads.route('/exportCsv')
def export_csv():
    """Makes a CSV file containing all information about each job."""

    message = helper_download('CSV')
    return jsonify( followup = message)




@downloads.route('/exportTxt')
def export_txt():
    """Makes a TXT file containing all information about each job."""

    message = helper_download('TXT')
    return jsonify( followup = message)



@downloads.route('/exportJson')
def export_json():
    """Creates a JSON file containing all information about each job."""
    """Modified to instead to create word document."""

    message = create_word_document()
    #message = helper_download('JSON')
    return jsonify( followup = message)



def helper_download(format, offset=0):
    """Since almost each download uses same code, this method reduces redundancy. If format is HTML, returns the data
    If the format is CSV, TXT, or JSON, returns the display message"""

    download = logic.Download()

    if format == 'HTML':
        return download.make_format(format, offset)

    if download.make_format(format):
        return format + ' file was successfully created!'
    else:
        return 'There was an error creating the ' + format + ' file. If it is already open, try closing and downloading it again'

def create_word_document():
    """Creates a word document containing job data"""

    document = Document()

    style=document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)


    def add_line(label, value):
        """Adds the paragraph with format of 'label: value[bolded]'"""
        para = document.add_paragraph()
        para_label = para.add_run(label + ": ")
        para_value = para.add_run(value)
        para_value.bold = True

    jobs = db.get_csv_json_text_data()

    for job in jobs:
        title = document.add_paragraph(job[1], style='Title')
        title.alignment = 1

        employer_name = document.add_paragraph(job[2], style='Subtitle')
        employer_name.alignment = 1

        add_line('Location', job[4])
        add_line('Level', job[6])
        add_line('Number of Openings', job[5])
        add_line('Discipline', job[7])
        add_line('Job Id', job[0])


        if (len(job[10]) > 10): #show Comments header only if the job contains comments
            document.add_heading('Comments', level=1)
            document.add_paragraph(job[10])

        # document.add_paragraph('\n\n\n\n\n\n\n')

        document.add_heading('Summary', level=1)
        #every time a <br /> occurs, add a linebreak to the current document
        print('The type of job[11] is: ' + str(type(job[11])))
        summary = job[11].replace('<br />', '\n').replace('&#039;', "'").replace('&nbsp;', ' ')
        document.add_paragraph(summary)

        document.add_page_break()

    document.save('jobs_EZSearch.docx')



